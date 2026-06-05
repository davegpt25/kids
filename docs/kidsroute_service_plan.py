from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── 색상 상수 ──
ORANGE      = RGBColor(0xF9, 0x73, 0x16)
ORANGE_DARK = RGBColor(0xEA, 0x6C, 0x0A)
DARK        = RGBColor(0x1C, 0x1C, 0x1C)
GRAY        = RGBColor(0x78, 0x71, 0x6C)
LIGHT_BG    = RGBColor(0xFF, 0xF7, 0xED)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
GREEN       = RGBColor(0x10, 0xB9, 0x81)
RED         = RGBColor(0xEF, 0x44, 0x44)
TABLE_HEADER= RGBColor(0xF9, 0x73, 0x16)
TABLE_ALT   = RGBColor(0xFF, 0xFB, 0xF7)

doc = Document()

# ── 기본 여백 설정 ──
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(3.0)

# ── 스타일 헬퍼 ──
def set_run_font(run, size, bold=False, color=None, italic=False):
    run.font.name  = '맑은 고딕'
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    rPr.insert(0, rFonts)

def set_para_spacing(para, before=0, after=6, line=None):
    pPr = para._p.get_or_add_pPr()
    pSpacing = OxmlElement('w:spacing')
    pSpacing.set(qn('w:before'), str(before))
    pSpacing.set(qn('w:after'),  str(after))
    if line:
        pSpacing.set(qn('w:line'),     str(line))
        pSpacing.set(qn('w:lineRule'), 'auto')
    pPr.append(pSpacing)

def add_heading(doc, text, level=1, color=DARK, size=None):
    sizes = {1: 22, 2: 16, 3: 13}
    para = doc.add_paragraph()
    run  = para.add_run(text)
    sz   = size or sizes.get(level, 12)
    set_run_font(run, sz, bold=True, color=color)
    set_para_spacing(para, before=240 if level==1 else 160, after=80)
    return para

def add_body(doc, text, indent=False, color=DARK, size=10.5):
    para = doc.add_paragraph()
    if indent:
        para.paragraph_format.left_indent = Cm(0.5)
    run  = para.add_run(text)
    set_run_font(run, size, color=color)
    set_para_spacing(para, before=0, after=40, line=280)
    return para

def add_bullet(doc, text, color=DARK, highlight=None):
    para = doc.add_paragraph(style='List Bullet')
    run  = para.add_run(text)
    set_run_font(run, 10, color=color)
    set_para_spacing(para, before=0, after=30)
    return para

def shade_cell(cell, color_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  color_hex)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, size=9.5, color=DARK, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    para = cell.paragraphs[0]
    para.alignment = align
    run  = para.add_run(text)
    set_run_font(run, size, bold=bold, color=color)
    set_para_spacing(para, before=30, after=30)

def add_divider(doc):
    para = doc.add_paragraph()
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),  'single')
    bottom.set(qn('w:sz'),   '4')
    bottom.set(qn('w:space'),'1')
    bottom.set(qn('w:color'),'F5F0EB')
    pBdr.append(bottom)
    pPr.append(pBdr)
    set_para_spacing(para, before=60, after=60)
    return para

def add_orange_box(doc, title, body_lines):
    """주황 배경 강조 박스"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, 'FFF7ED')
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        bdr = OxmlElement(f'w:{side}')
        bdr.set(qn('w:val'),   'single')
        bdr.set(qn('w:sz'),    '12')
        bdr.set(qn('w:space'), '0')
        bdr.set(qn('w:color'), 'F97316')
        tcBorders.append(bdr)
    tcPr.append(tcBorders)
    cell.text = ''
    p1 = cell.add_paragraph()
    r1 = p1.add_run(title)
    set_run_font(r1, 11, bold=True, color=ORANGE)
    set_para_spacing(p1, before=60, after=40)
    for line in body_lines:
        p2 = cell.add_paragraph()
        r2 = p2.add_run(line)
        set_run_font(r2, 10, color=DARK)
        set_para_spacing(p2, before=0, after=30)
    doc.add_paragraph()  # spacing after box

# ════════════════════════════════════════════
# 표지
# ════════════════════════════════════════════
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(cover, before=600, after=120)
r = cover.add_run('키즈루트 (KidsRoute)')
set_run_font(r, 28, bold=True, color=ORANGE)

cover2 = doc.add_paragraph()
cover2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = cover2.add_run('서비스 기획서')
set_run_font(r2, 18, bold=True, color=DARK)
set_para_spacing(cover2, before=0, after=80)

cover3 = doc.add_paragraph()
cover3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = cover3.add_run('3초면 완성되는 우리 아이 학원 스케줄')
set_run_font(r3, 13, italic=True, color=GRAY)
set_para_spacing(cover3, before=0, after=60)

cover4 = doc.add_paragraph()
cover4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = cover4.add_run('MVP v1.0 · 서울 강남구 출시 예정  |  2026년 5월')
set_run_font(r4, 10, color=GRAY)
set_para_spacing(cover4, before=0, after=600)

doc.add_page_break()

# ════════════════════════════════════════════
# 1. 서비스 개요
# ════════════════════════════════════════════
add_heading(doc, '1. 서비스 개요', level=1, color=ORANGE)
add_divider(doc)

add_heading(doc, '1.1 서비스 정의', level=2)
add_body(doc,
    '키즈루트는 학부모를 위한 학원 스케줄 최적화 플랫폼이다. '
    'GPS 기반으로 반경 내 학원을 자동 수집하고, 시간 충돌 없는 최적 조합을 3초 만에 추천한다. '
    '나아가 같은 방향 이웃 아이들을 연결해 학원 주도의 셔틀 정류장 형성을 지원한다.')

t = doc.add_table(rows=5, cols=2)
t.style = 'Table Grid'
rows_data = [
    ('서비스명',  '키즈루트 (KidsRoute)'),
    ('태그라인',  '3초면 완성되는 우리 아이 학원 스케줄'),
    ('플랫폼',   'iOS (App Store) · Android (Play Store)'),
    ('1차 출시',  '서울 강남구 MVP → 전국 확장'),
    ('현재 단계', '사전예약 랜딩페이지 운영 중 (davegpt25.github.io/kids)'),
]
for i, (k, v) in enumerate(rows_data):
    row = t.rows[i]
    shade_cell(row.cells[0], 'FFF7ED')
    set_cell_text(row.cells[0], k, bold=True, size=9.5, color=ORANGE)
    set_cell_text(row.cells[1], v, size=9.5)
doc.add_paragraph()

add_heading(doc, '1.2 핵심 가치 제안 (Value Proposition)', level=2)
bullets = [
    '학부모의 학원 탐색·조합·관리 시간을 주 평균 3시간 → 3분으로 단축',
    '이동 시간 포함 시간 충돌 자동 감지로 학원 중복 문제 해소',
    '5명이 같은 방향에 모이면 학원이 직접 셔틀 정류장을 만들 수 있어 교통 비용 절감',
    '3명 이상 동시 하원 시 하원 서비스 무료 — 네트워크 효과로 자연 성장',
]
for b in bullets:
    add_bullet(doc, b)
doc.add_paragraph()

add_heading(doc, '1.3 문제 정의', level=2)
prob_table = doc.add_table(rows=4, cols=2)
prob_table.style = 'Table Grid'
header_row = prob_table.rows[0]
shade_cell(header_row.cells[0], 'F97316')
shade_cell(header_row.cells[1], 'F97316')
set_cell_text(header_row.cells[0], '현재 문제 (As-Is)', bold=True, color=WHITE, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_text(header_row.cells[1], '키즈루트 해결 (To-Be)', bold=True, color=WHITE, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
problems = [
    ('학원 정보가 흩어져 있어 직접 발품 팔거나 맘카페 의존', 'NEIS 공공데이터 기반 GPS 반경 학원 자동 수집'),
    ('시간표 충돌 여부를 엑셀·수기로 직접 확인', '이동 10분 버퍼 포함 충돌 자동 감지 & 최적 조합 추천'),
    ('등하원 셔틀 없는 학원 = 학부모 픽업 부담', '방향별 수요 집계 → 5명 충족 시 학원이 셔틀 직접 운영'),
]
for i, (p, s) in enumerate(problems, 1):
    row = prob_table.rows[i]
    if i % 2 == 0:
        shade_cell(row.cells[0], 'FFF7ED')
        shade_cell(row.cells[1], 'FFF7ED')
    set_cell_text(row.cells[0], p, size=9)
    set_cell_text(row.cells[1], s, size=9)
doc.add_paragraph()

doc.add_page_break()

# ════════════════════════════════════════════
# 2. 타겟 사용자
# ════════════════════════════════════════════
add_heading(doc, '2. 타겟 사용자', level=1, color=ORANGE)
add_divider(doc)

add_heading(doc, '2.1 주 사용자 — 학부모', level=2)
user_tbl = doc.add_table(rows=5, cols=2)
user_tbl.style = 'Table Grid'
user_data = [
    ('연령대',  '30~40대 자녀 보유 부모 (초등학교 1~6학년 자녀 기준)'),
    ('거주지',  '서울 강남·서초·송파구 → 전국 확장'),
    ('Pain Point', '학원 조합 탐색·시간 충돌 확인·등하원 교통 문제'),
    ('사용 계기', '맘카페·지인 추천, 앱스토어 검색'),
    ('결정 요인', '편의성·신뢰성·비용 절감'),
]
for i, (k, v) in enumerate(user_data):
    row = user_tbl.rows[i]
    shade_cell(row.cells[0], 'FFF7ED')
    set_cell_text(row.cells[0], k, bold=True, size=9.5, color=ORANGE)
    set_cell_text(row.cells[1], v, size=9.5)
doc.add_paragraph()

add_heading(doc, '2.2 B2B 파트너 — 학원', level=2)
add_body(doc,
    '학원은 키즈루트의 B2B 파트너로, 셔틀 정류장 운영 주체다. '
    '학원이 방향별 탑승 수요(5명 기준)를 확인하고 직접 셔틀을 운행하며, '
    '키즈루트는 수요 집계 및 알림 인프라를 제공한다.')
b2b = [
    '학원 입장: 기존 통학차 노선 최적화 + 신규 수요 발굴 가능',
    '유치원 픽업: 5명 모집 시 특정 유치원 픽업 서비스 개설 (학원 주도)',
    '향후 프리미엄 B2B 요금제로 학원 대상 수익화 가능',
]
for b in b2b:
    add_bullet(doc, b)
doc.add_paragraph()

doc.add_page_break()

# ════════════════════════════════════════════
# 3. 핵심 기능
# ════════════════════════════════════════════
add_heading(doc, '3. 핵심 기능 (As-Is MVP)', level=1, color=ORANGE)
add_divider(doc)

features = [
    ('📍 GPS 반경 학원 탐색',
     '거주지·학교 기준 500m~5km 반경 학원 자동 수집. NEIS 공공데이터 연동. '
     '카카오맵 API 기반 실시간 지도 시각화.'),
    ('⚡ 시간 충돌 자동 감지',
     '이동 시간 10분 버퍼 포함 학원 간 시간 충돌 자동 검사. '
     'Greedy 알고리즘으로 충돌 없는 최적 조합 순위 추천.'),
    ('🏆 과목 우선순위 조합',
     '수학·영어·예체능 등 과목 우선순위 설정. 조합 수 : 무료(1개) / 프리미엄(무제한).'),
    ('📅 주간 캘린더 시각화',
     '추천 조합을 월~금 블록 타임라인으로 표시. 아이 프로필별 독립 캘린더.'),
    ('🔐 카카오 / 구글 소셜 로그인',
     '팝업(데스크톱) + 리다이렉트(모바일) 방식 OAuth. JWT 기반 인증. '
     '모바일 팝업 차단 이슈 해결 완료.'),
    ('🚌 학원 주도 셔틀 정류장',
     '방향별 탑승 수요 실시간 집계. 5명 충족 시 학원이 정류장 개설·운행. '
     '학부모 탭·학원 운영 탭 분리 UI.'),
    ('🏠 하원 서비스',
     '3명 이상 동시 하원 시 무료. 키즈루트가 자동 매칭. '
     '네트워크 효과로 이용자 증가 시 비용 자동 절감.'),
    ('💾 스케줄 저장 & 관리',
     '마음에 드는 조합 저장. 복수 자녀 프로필 독립 관리(프리미엄).'),
]

feat_tbl = doc.add_table(rows=len(features)+1, cols=2)
feat_tbl.style = 'Table Grid'
hrow = feat_tbl.rows[0]
shade_cell(hrow.cells[0], 'F97316')
shade_cell(hrow.cells[1], 'F97316')
set_cell_text(hrow.cells[0], '기능', bold=True, color=WHITE, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_text(hrow.cells[1], '설명', bold=True, color=WHITE, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
for i, (name, desc) in enumerate(features, 1):
    row = feat_tbl.rows[i]
    if i % 2 == 0:
        shade_cell(row.cells[0], 'FFF7ED')
        shade_cell(row.cells[1], 'FFF7ED')
    set_cell_text(row.cells[0], name, bold=True, size=9.5)
    set_cell_text(row.cells[1], desc, size=9)

# col widths
feat_tbl.columns[0].width = Cm(5)
feat_tbl.columns[1].width = Cm(11)
doc.add_paragraph()

doc.add_page_break()

# ════════════════════════════════════════════
# 4. 수익 모델 (핵심)
# ════════════════════════════════════════════
add_heading(doc, '4. 수익 모델', level=1, color=ORANGE)
add_divider(doc)

add_orange_box(doc,
    '💡 수익 전략 요약',
    [
        '① B2C 구독 (학부모 프리미엄) — 핵심 초기 수익원',
        '② 셔틀 네트워크 수수료 (학원 B2B) — 중기 수익원',
        '③ 학원 광고·추천 수익 — 장기 수익원',
        '④ 데이터 인사이트 판매 — 장기 수익원',
    ]
)

# 4.1 B2C 프리미엄 구독
add_heading(doc, '4.1 B2C 프리미엄 구독 (학부모)', level=2)

plan_tbl = doc.add_table(rows=10, cols=3)
plan_tbl.style = 'Table Grid'
plan_headers = ['항목', '무료 플랜', '프리미엄 플랜']
plan_shades  = ['F97316', 'F97316', 'F97316']
for j, (h, sh) in enumerate(zip(plan_headers, plan_shades)):
    shade_cell(plan_tbl.rows[0].cells[j], sh)
    set_cell_text(plan_tbl.rows[0].cells[j], h, bold=True, color=WHITE, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)

plan_rows = [
    ('월 요금',             '₩0 (영원히 무료)',   '₩9,900 / 월'),
    ('GPS 반경 학원 탐색',  '✓',                '✓'),
    ('스케줄 조합 추천',    '1개',              '무제한'),
    ('시간 충돌 알림',      '✗',                '✓'),
    ('복수 자녀 관리',      '1명',              '무제한'),
    ('예산 필터',          '✗',                '✓'),
    ('학원 시간표 변경 알림','✗',                '✓'),
    ('우선 고객 지원',      '✗',                '✓'),
    ('출시 시기',          'Phase 1 (즉시)',    'Phase 2 (4개월차~)'),
]
for i, (item, free, paid) in enumerate(plan_rows, 1):
    row = plan_tbl.rows[i]
    if i % 2 == 0:
        shade_cell(row.cells[0], 'FFF7ED')
        shade_cell(row.cells[1], 'FFF7ED')
        shade_cell(row.cells[2], 'FFF7ED')
    set_cell_text(row.cells[0], item, bold=(i==1), size=9.5)
    set_cell_text(row.cells[1], free, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER,
                  color=(GRAY if free=='✗' else GREEN if free=='✓' else DARK))
    set_cell_text(row.cells[2], paid, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER,
                  color=(GREEN if paid=='✓' else ORANGE if paid.startswith('₩') else DARK))

plan_tbl.columns[0].width = Cm(5)
plan_tbl.columns[1].width = Cm(5)
plan_tbl.columns[2].width = Cm(6)
doc.add_paragraph()

add_body(doc,
    '▶ 수익 시뮬레이션 (Phase 2 기준: MAU 5,000명, 전환율 10%)',
    color=GRAY, size=9.5)
rev_sim = doc.add_table(rows=4, cols=3)
rev_sim.style = 'Table Grid'
for j, h in enumerate(['구분', '수치', '비고']):
    shade_cell(rev_sim.rows[0].cells[j], 'F97316')
    set_cell_text(rev_sim.rows[0].cells[j], h, bold=True, color=WHITE, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
sim_rows = [
    ('프리미엄 유저 수',     '500명',          'MAU 5,000명 × 전환율 10%'),
    ('월 구독 단가',         '₩9,900',         '경쟁사 대비 합리적 가격'),
    ('월 구독 매출 (MRR)',   '₩4,950,000',     '약 495만원 / 월'),
]
for i, (a, b, c) in enumerate(sim_rows, 1):
    row = rev_sim.rows[i]
    if i % 2 == 0:
        for cell in row.cells: shade_cell(cell, 'FFF7ED')
    set_cell_text(row.cells[0], a, size=9.5)
    set_cell_text(row.cells[1], b, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row.cells[2], c, size=9)
doc.add_paragraph()

# 4.2 셔틀 네트워크 수수료
add_heading(doc, '4.2 셔틀 네트워크 수수료 (학원 B2B)', level=2)
add_body(doc,
    '학원이 키즈루트 플랫폼을 통해 셔틀 정류장을 개설하고 탑승 수요를 관리할 때 수수료를 부과한다. '
    '핵심 인프라는 키즈루트가 제공하므로 학원은 별도 앱 개발 없이 셔틀 운영이 가능하다.')

shuttle_models = [
    ('기본 수요 집계 알림',  '무료',            '학원 유입 확대 목적'),
    ('프리미엄 대시보드',    '₩29,900 / 월',   '실시간 탑승 수요·노선 관리'),
    ('다중 노선 관리',       '₩59,900 / 월',   '노선 5개 이상 학원 대상'),
    ('유치원 픽업 연결',     '건당 수수료 5%',  '픽업 서비스 예약 당 수수료'),
]
sh_tbl = doc.add_table(rows=len(shuttle_models)+1, cols=3)
sh_tbl.style = 'Table Grid'
for j, h in enumerate(['서비스', '요금', '비고']):
    shade_cell(sh_tbl.rows[0].cells[j], 'F97316')
    set_cell_text(sh_tbl.rows[0].cells[j], h, bold=True, color=WHITE, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
for i, (s, p, n) in enumerate(shuttle_models, 1):
    row = sh_tbl.rows[i]
    if i % 2 == 0:
        for cell in row.cells: shade_cell(cell, 'FFF7ED')
    set_cell_text(row.cells[0], s, size=9.5)
    set_cell_text(row.cells[1], p, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row.cells[2], n, size=9)
doc.add_paragraph()

# 4.3 학원 광고·추천 수익
add_heading(doc, '4.3 학원 광고 & 추천 노출 수익', level=2)
add_body(doc,
    '키즈루트 검색 결과 상단 노출 광고 상품. '
    '학원이 특정 과목·지역에서 우선 노출을 구매하는 방식으로, '
    'Google 지역 검색 광고 모델과 유사하다.')
ad_bullets = [
    '상단 노출 광고: 주 단위 경매 방식, 예상 단가 ₩50,000~₩300,000/주',
    '추천 배너 (학원 이벤트·할인): 클릭당 과금(CPC) 방식',
    '학원 프로필 인증 배지: ₩9,900/월 — 신뢰도 강조',
]
for b in ad_bullets:
    add_bullet(doc, b)
doc.add_paragraph()

# 4.4 데이터 인사이트
add_heading(doc, '4.4 데이터 인사이트 판매 (장기)', level=2)
add_body(doc,
    '학원 수요·이동 패턴·과목 트렌드 분석 데이터를 집계·비식별화하여 '
    '교육 기관, 부동산 개발사, 지자체 등에 제공한다.')
data_bullets = [
    '교육 수요 히트맵: 과목별·지역별 수요 밀도 리포트',
    '과목 트렌드 리포트: 분기별 인기 과목·학원 분석',
    '입주 개발 컨설팅: 신축 아파트 단지 학원 수요 예측 리포트',
]
for b in data_bullets:
    add_bullet(doc, b)
doc.add_paragraph()

# 4.5 수익 단계별 로드맵
add_heading(doc, '4.5 수익 단계별 로드맵', level=2)
phase_tbl = doc.add_table(rows=4, cols=4)
phase_tbl.style = 'Table Grid'
ph_headers = ['단계', '기간', '수익원', '목표 MRR']
for j, h in enumerate(ph_headers):
    shade_cell(phase_tbl.rows[0].cells[j], 'F97316')
    set_cell_text(phase_tbl.rows[0].cells[j], h, bold=True, color=WHITE, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
phases = [
    ('Phase 1', '0~3개월',  'MAU 500명 / 무료 플랜 위주 / 사전예약 전환',    '-'),
    ('Phase 2', '4~6개월',  '프리미엄 구독 ₩9,900 출시 / 전환율 10% 목표', '₩495만'),
    ('Phase 3', '7~12개월', '학원 B2B 요금제 + 광고 수익 / MAU 30,000명',  '₩5,000만'),
]
for i, row_data in enumerate(phases, 1):
    row = phase_tbl.rows[i]
    if i % 2 == 0:
        for cell in row.cells: shade_cell(cell, 'FFF7ED')
    for j, val in enumerate(row_data):
        bold = (j == 0)
        align = WD_ALIGN_PARAGRAPH.CENTER if j in [0, 3] else WD_ALIGN_PARAGRAPH.LEFT
        color = ORANGE if j == 0 else DARK
        set_cell_text(row.cells[j], val, bold=bold, size=9, color=color, align=align)
doc.add_paragraph()

doc.add_page_break()

# ════════════════════════════════════════════
# 5. 사용자 여정 (User Journey)
# ════════════════════════════════════════════
add_heading(doc, '5. 사용자 여정 (User Journey)', level=1, color=ORANGE)
add_divider(doc)

add_heading(doc, '5.1 학부모 온보딩 플로우', level=2)
journey = [
    ('1', '앱 설치', 'App Store / Play Store 검색 → 설치'),
    ('2', '소셜 로그인', 'Google 또는 카카오 계정으로 1-tap 로그인 (팝업/리다이렉트)'),
    ('3', '위치 설정', 'GPS 자동 감지 또는 주소 입력 → 반경(500m~5km) 선택'),
    ('4', '과목 선택', '원하는 과목 및 우선순위 설정'),
    ('5', '조합 추천', '충돌 없는 최적 학원 조합 3초 내 표시'),
    ('6', '캘린더 확인', '주간 타임라인으로 스케줄 시각 확인'),
    ('7', '저장 & 공유', '조합 저장 후 가족과 공유'),
    ('8', '셔틀 신청', '원하는 방향 셔틀 탑승 신청 → 5명 모이면 운행 시작'),
]
jrn_tbl = doc.add_table(rows=len(journey)+1, cols=3)
jrn_tbl.style = 'Table Grid'
for j, h in enumerate(['단계', '액션', '설명']):
    shade_cell(jrn_tbl.rows[0].cells[j], 'F97316')
    set_cell_text(jrn_tbl.rows[0].cells[j], h, bold=True, color=WHITE, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
for i, (step, action, desc) in enumerate(journey, 1):
    row = jrn_tbl.rows[i]
    if i % 2 == 0:
        for cell in row.cells: shade_cell(cell, 'FFF7ED')
    set_cell_text(row.cells[0], step, bold=True, size=9.5, color=ORANGE, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row.cells[1], action, bold=True, size=9.5)
    set_cell_text(row.cells[2], desc, size=9)
jrn_tbl.columns[0].width = Cm(1.5)
jrn_tbl.columns[1].width = Cm(4)
jrn_tbl.columns[2].width = Cm(10.5)
doc.add_paragraph()

add_heading(doc, '5.2 데모 체험 플로우 (랜딩페이지)', level=2)
demo_steps = [
    ('Step 1', '타이프라이터', '예시 검색어 5종 순환 애니메이션 → 사용자가 원하는 쿼리 선택'),
    ('Step 2', '지도 설정',   '카카오맵 GPS 위치 표시 + 반경 선택 (500m~5km) + 셔틀/하원 서비스 토글'),
    ('Step 3', '분석 중',     '프로그레스 바 0→100% 애니메이션 (~2.5초)'),
    ('Step 4', '리드 폼',     '이메일·성함 입력 또는 Google/카카오로 자동 입력'),
    ('Step 5', '완료',        '사전예약 확인 + 얼리버드 3개월 프리미엄 무료 혜택 안내'),
]
for step, name, desc in demo_steps:
    p = doc.add_paragraph()
    r1 = p.add_run(f'{step} — {name}  ')
    set_run_font(r1, 10, bold=True, color=ORANGE)
    r2 = p.add_run(desc)
    set_run_font(r2, 10, color=DARK)
    set_para_spacing(p, before=0, after=40)
doc.add_paragraph()

doc.add_page_break()

# ════════════════════════════════════════════
# 6. 기술 스택
# ════════════════════════════════════════════
add_heading(doc, '6. 기술 스택 (현재 MVP)', level=1, color=ORANGE)
add_divider(doc)

tech_tbl = doc.add_table(rows=9, cols=2)
tech_tbl.style = 'Table Grid'
hrow = tech_tbl.rows[0]
shade_cell(hrow.cells[0], 'F97316')
shade_cell(hrow.cells[1], 'F97316')
set_cell_text(hrow.cells[0], '구분', bold=True, color=WHITE, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_text(hrow.cells[1], '기술', bold=True, color=WHITE, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
tech_data = [
    ('프론트엔드',  'HTML5 · CSS3 · Vanilla JS (단일 index.html, 라이브러리 무사용)'),
    ('지도 API',   'Kakao Maps JavaScript API (무료, 기존 앱키 재사용)'),
    ('인증 (Google)', 'Google Identity Services — Token Client (데스크톱 팝업 + 모바일 implicit redirect)'),
    ('인증 (Kakao)',  'Kakao JS SDK v2.7.2 — Auth.login 팝업 + Authorization Code redirect (모바일)'),
    ('공공데이터',  'NEIS (학원 정보) API 연동 예정'),
    ('배포',       'GitHub Pages (github.com/davegpt25/kids → gh-pages 브랜치)'),
    ('분석 예정',   'Google Analytics 4 · Firebase A/B Testing'),
    ('백엔드 예정', 'Node.js + PostgreSQL (Phase 2 사전예약 DB 저장)'),
]
for i, (k, v) in enumerate(tech_data, 1):
    row = tech_tbl.rows[i]
    if i % 2 == 0:
        shade_cell(row.cells[0], 'FFF7ED')
        shade_cell(row.cells[1], 'FFF7ED')
    set_cell_text(row.cells[0], k, bold=True, size=9.5, color=ORANGE)
    set_cell_text(row.cells[1], v, size=9)
tech_tbl.columns[0].width = Cm(4)
tech_tbl.columns[1].width = Cm(12)
doc.add_paragraph()

doc.add_page_break()

# ════════════════════════════════════════════
# 7. KPI & 성장 목표
# ════════════════════════════════════════════
add_heading(doc, '7. KPI & 성장 목표', level=1, color=ORANGE)
add_divider(doc)

kpi_tbl = doc.add_table(rows=4, cols=5)
kpi_tbl.style = 'Table Grid'
kpi_headers = ['지표', 'Phase 1 (0~3M)', 'Phase 2 (4~6M)', 'Phase 3 (7~12M)', '측정 방법']
for j, h in enumerate(kpi_headers):
    shade_cell(kpi_tbl.rows[0].cells[j], 'F97316')
    set_cell_text(kpi_tbl.rows[0].cells[j], h, bold=True, color=WHITE, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
kpi_data = [
    ('MAU',        '500명',    '5,000명',    '30,000명',  'Firebase Analytics'),
    ('프리미엄 전환율', '-', '10%', '15%+', '결제 시스템 집계'),
    ('MRR',        '-',       '₩495만',     '₩5,000만',  '구독 매출 합산'),
]
for i, row_d in enumerate(kpi_data, 1):
    row = kpi_tbl.rows[i]
    if i % 2 == 0:
        for cell in row.cells: shade_cell(cell, 'FFF7ED')
    for j, val in enumerate(row_d):
        align = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
        bold = (j == 0)
        set_cell_text(row.cells[j], val, bold=bold, size=9, align=align)
doc.add_paragraph()

add_heading(doc, '7.2 사전예약 전환 목표', level=2)
add_bullet(doc, '랜딩페이지 방문 → 사전예약 전환율 목표: 15% 이상')
add_bullet(doc, '사전예약 → 앱 설치 전환율 목표: 50% 이상 (푸시 알림 기반)')
add_bullet(doc, '얼리버드 혜택(3개월 프리미엄 무료)으로 초기 프리미엄 사용자 확보')
doc.add_paragraph()

doc.add_page_break()

# ════════════════════════════════════════════
# 8. 경쟁 분석
# ════════════════════════════════════════════
add_heading(doc, '8. 경쟁 분석', level=1, color=ORANGE)
add_divider(doc)

comp_tbl = doc.add_table(rows=5, cols=5)
comp_tbl.style = 'Table Grid'
comp_headers = ['서비스', '학원 탐색', '시간 최적화', '셔틀 연계', '가격']
for j, h in enumerate(comp_headers):
    shade_cell(comp_tbl.rows[0].cells[j], 'F97316')
    set_cell_text(comp_tbl.rows[0].cells[j], h, bold=True, color=WHITE, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
comp_data = [
    ('키즈루트 (자사)',   '✓ GPS 자동', '✓ 충돌 자동 감지', '✓ 학원 주도',  '₩0 / ₩9,900'),
    ('네이버 지도',      '✓ 검색',     '✗ 수동 확인',      '✗',           '무료'),
    ('학원 자체 앱',     '△ 해당 학원만', '✗',             '△ 학원별',    '무료'),
    ('에듀픽 / 클래스팅','✓ 정보 제공', '✗',              '✗',           '무료/유료'),
]
for i, row_d in enumerate(comp_data, 1):
    row = comp_tbl.rows[i]
    if i % 2 == 0:
        for cell in row.cells: shade_cell(cell, 'FFF7ED')
    for j, val in enumerate(row_d):
        bold = (i == 1)  # 자사 행 bold
        color = ORANGE if (i == 1 and j == 0) else DARK
        set_cell_text(row.cells[j], val, bold=bold, size=9, color=color, align=WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT)
doc.add_paragraph()

add_body(doc,
    '키즈루트의 핵심 차별점: ① 학원 조합 시간 충돌 자동 감지, '
    '② 학원 주도 셔틀 정류장 형성 메커니즘, '
    '③ 하원 서비스(3명 이상 동시 하원 무료) — 세 가지 모두 기존 경쟁사에 없는 기능이다.',
    color=GRAY)
doc.add_paragraph()

doc.add_page_break()

# ════════════════════════════════════════════
# 9. 위험 요소 & 대응 전략
# ════════════════════════════════════════════
add_heading(doc, '9. 위험 요소 & 대응 전략', level=1, color=ORANGE)
add_divider(doc)

risk_tbl = doc.add_table(rows=5, cols=3)
risk_tbl.style = 'Table Grid'
for j, h in enumerate(['위험 요소', '영향도', '대응 전략']):
    shade_cell(risk_tbl.rows[0].cells[j], 'F97316')
    set_cell_text(risk_tbl.rows[0].cells[j], h, bold=True, color=WHITE, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
risks = [
    ('초기 사용자 확보 어려움 (콜드 스타트)',       '높음', '강남구 맘카페·학부모 커뮤니티 집중 마케팅 + 얼리버드 혜택'),
    ('학원 데이터 정확도 (NEIS 갱신 지연)',       '중간', 'NEIS 외 학원 자체 신고 채널 병행, 커뮤니티 기반 제보 시스템'),
    ('셔틀 5명 미달 시 서비스 미작동 (네트워크 효과)', '높음', '지역별 수요 집중 전략 (강남구 한 지역 먼저 밀도 확보)'),
    ('Google/Kakao OAuth 정책 변경',          '낮음', '리다이렉트 방식으로 이미 대응 완료, SDK 버전 고정 관리'),
]
for i, (r, impact, action) in enumerate(risks, 1):
    row = risk_tbl.rows[i]
    if i % 2 == 0:
        for cell in row.cells: shade_cell(cell, 'FFF7ED')
    imp_color = RED if impact == '높음' else ORANGE if impact == '중간' else GREEN
    set_cell_text(row.cells[0], r, size=9)
    set_cell_text(row.cells[1], impact, size=9.5, bold=True, color=imp_color, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row.cells[2], action, size=9)
risk_tbl.columns[0].width = Cm(5)
risk_tbl.columns[1].width = Cm(1.5)
risk_tbl.columns[2].width = Cm(9.5)
doc.add_paragraph()

doc.add_page_break()

# ════════════════════════════════════════════
# 10. 출시 로드맵
# ════════════════════════════════════════════
add_heading(doc, '10. 출시 로드맵', level=1, color=ORANGE)
add_divider(doc)

road_tbl = doc.add_table(rows=4, cols=3)
road_tbl.style = 'Table Grid'
for j, h in enumerate(['단계', '주요 과업', '성과 지표']):
    shade_cell(road_tbl.rows[0].cells[j], 'F97316')
    set_cell_text(road_tbl.rows[0].cells[j], h, bold=True, color=WHITE, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
roadmap = [
    ('Phase 1\n(0~3개월)',
     '• 강남구 MVP 앱 출시\n• 사전예약 → 앱 설치 전환 캠페인\n• 무료 플랜 위주 운영\n• 셔틀 수요 데이터 수집',
     'MAU 500명\n사전예약 500건'),
    ('Phase 2\n(4~6개월)',
     '• 프리미엄 구독 (₩9,900/월) 출시\n• 학원 B2B 대시보드 베타\n• 서초·송파구 확장\n• 광고 상품 출시',
     'MAU 5,000명\nMRR ₩495만\n전환율 10%'),
    ('Phase 3\n(7~12개월)',
     '• 전국 확장 (서울 전역 → 수도권)\n• 데이터 인사이트 상품 출시\n• 유치원 픽업 연계 서비스 정식화\n• Series A 투자 준비',
     'MAU 30,000명\nMRR ₩5,000만'),
]
for i, (phase, tasks, kpi) in enumerate(roadmap, 1):
    row = road_tbl.rows[i]
    if i % 2 == 0:
        for cell in row.cells: shade_cell(cell, 'FFF7ED')
    set_cell_text(row.cells[0], phase, bold=True, size=9.5, color=ORANGE, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row.cells[1], tasks, size=9)
    set_cell_text(row.cells[2], kpi, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
road_tbl.columns[0].width = Cm(2.5)
road_tbl.columns[1].width = Cm(10)
road_tbl.columns[2].width = Cm(3.5)
doc.add_paragraph()

# ════════════════════════════════════════════
# 마무리
# ════════════════════════════════════════════
add_divider(doc)
end_para = doc.add_paragraph()
end_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_end = end_para.add_run('키즈루트 — 우리 아이 학원 스케줄, 이제 3초면 완성됩니다.')
set_run_font(r_end, 12, bold=True, color=ORANGE)
set_para_spacing(end_para, before=200, after=200)

out = r'C:\Users\hwlll\Startup\kids\docs\키즈루트_서비스기획서_v1.docx'
doc.save(out)
print('saved:', out)
